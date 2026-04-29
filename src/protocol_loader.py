"""
Protocol loader module.

Downloads a protocol .docx from Google Drive, extracts all text (body paragraphs
+ table cells), and loads the companion knowledge Google Doc if one exists.

Returns the text components separately so that build_system_prompt() in
claude_client.py can format them into the Protocol Expert system prompt.

IMPORTANT — why we extract tables separately:
    doc.paragraphs from python-docx only covers body text.
    Table cells are NOT included. Buffer recipes almost always live in tables,
    so iterating doc.tables is critical for the /buffer functionality to work.
"""

from __future__ import annotations

import io
import logging
import os
from typing import Optional

from docx import Document

from .google_client import download_docx, find_companion_doc_id, get_doc_text

logger = logging.getLogger(__name__)


def extract_docx_text(docx_bytes: bytes) -> str:
    """Extract all text from a .docx file — body paragraphs + table contents.

    Table rows are formatted as pipe-separated cells so Claude can read
    them as structured data (e.g. reagent | stock concentration | final volume).
    """
    doc = Document(io.BytesIO(docx_bytes))
    parts: list[str] = []

    # Body paragraphs (headings, steps, notes)
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Tables — critical for buffer recipes
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


async def load_protocol(
    file_id: str,
    file_name: str,
    modified_time: str,
    parent_folder_id: str = "",
    folder_name: str = "",
    is_gdoc: bool = False,
) -> tuple[str, str, str, str, Optional[str]]:
    """Load a protocol from Drive and return all components for the skill.

    Supports both .docx files and native Google Docs.
    """
    # Use folder name as protocol name when the file itself is just "method_support"
    raw_name = os.path.splitext(file_name)[0]
    protocol_name = folder_name if (folder_name and raw_name == "method_support") else raw_name
    mod_date = modified_time[:10] if modified_time else "unknown"
    protocol_version = f"{file_name} (modified {mod_date})"

    # Download and extract text
    if is_gdoc:
        protocol_text = await get_doc_text(file_id)
        logger.info("Loaded Google Doc protocol '%s' (%d chars)", protocol_name, len(protocol_text))
    else:
        docx_bytes = await download_docx(file_id)
        protocol_text = extract_docx_text(docx_bytes)
        logger.info("Loaded protocol '%s' (%d chars)", protocol_name, len(protocol_text))

    # Search companion using folder name (preferred) or file stem (fallback).
    # Skip companion search if the protocol file IS the method_support doc
    # (otherwise we'd double-load the same document).
    search_name = folder_name or protocol_name
    companion_text = ""
    companion_doc_id: Optional[str] = None
    _is_method_support = file_name.lower().replace(" ", "_") in ("method_support",)
    if _is_method_support:
        logger.info("Protocol IS method_support — skipping companion search")
        companion_doc_id = file_id  # keep ID for /refine appending
    try:
        if not _is_method_support:
            companion_doc_id = await find_companion_doc_id(search_name, parent_folder_id)
        if companion_doc_id and not _is_method_support:
            companion_text = await get_doc_text(companion_doc_id)
            logger.info(
                "Loaded companion doc for '%s' (%d chars)", protocol_name, len(companion_text)
            )
    except Exception:
        logger.warning(
            "Could not load companion doc for '%s'", protocol_name, exc_info=True
        )

    return protocol_text, companion_text, protocol_name, protocol_version, companion_doc_id
